import os
import re
import json
import tempfile
from pathlib import Path
from typing import List, Dict, Tuple

from docx import Document
import faiss
import glob
import pickle

# LangChain imports (langchain-community / langchain-openai)
from langchain_community.vectorstores import FAISS as LC_FAISS
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import HumanMessage

# ---------- CONFIG ----------
FAISS_DIR = os.environ.get("FAISS_DIR", "faiss_store") 
REPORTS_DIR = os.environ.get("REPORTS_DIR", "reports")
os.makedirs(REPORTS_DIR, exist_ok=True)

EMBEDDING_MODEL = "text-embedding-3-small"
LLM_MODEL_NAME = os.environ.get("LLM_MODEL_NAME", "gpt-4o-mini")

# ----------------------------

# simple document-type keywords
DOCUMENT_KEYWORDS = {
    "Articles_of_Association": ["articles of association", "aoa"],
    "Memorandum_of_Association": ["memorandum of association", "moa", "mou"],
    "Board_Resolution_Template": ["board resolution"],
    "Shareholder_Resolution_Template": ["shareholder resolution"],
    "Incorporation_Application_Form": ["incorporation application"],
    "UBO_Declaration_Form": ["ubo declaration", "ultimate beneficial"],
    "Register_of_Members_and_Directors": ["register of members", "register of directors"],
    "Change_of_Registered_Address_Notice": ["change of registered address"]
}

# Checklists for processes (extendable)
ADGM_CHECKLISTS = {
    "Company Incorporation": [
        "Articles_of_Association",
        "Memorandum_of_Association",
        "Board_Resolution_Template",
        "Shareholder_Resolution_Template",
        "Incorporation_Application_Form",
        "UBO_Declaration_Form",
        "Register_of_Members_and_Directors",
        "Change_of_Registered_Address_Notice"
    ],
    # we can other processes e.g., "Licensing": [...], "HR": [...]
}

# ---------- Utilities ----------

def extract_docx_text(path: str) -> str:
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paras)

def classify_doc_by_keywords(filename: str, text: str) -> str:
    fn = filename.lower()
    tx = text.lower()
    for dtype, kws in DOCUMENT_KEYWORDS.items():
        if any(k in fn for k in kws) or any(k in tx for k in kws):
            return dtype
    return "Unknown"

def detect_process(uploaded_doc_types: List[str]) -> Tuple[str, List[str]]:
    """
    Choose the process with highest overlap; return (process_name, missing_required_list)
    """
    best = None
    best_overlap = -1
    for proc, req in ADGM_CHECKLISTS.items():
        overlap = len(set(req) & set(uploaded_doc_types))
        if overlap > best_overlap:
            best_overlap = overlap
            best = proc
    if best is None:
        return "Unknown", []
    required = ADGM_CHECKLISTS[best]
    missing = [r for r in required if r not in uploaded_doc_types]
    return best, missing

# ---------- Load FAISS retriever & LLM ----------

def load_retriever_and_llm(faiss_dir: str = FAISS_DIR):
    """
    Load the saved FAISS vectorstore (assumes it was saved with LangChain's FAISS.from_documents(...).save_local)
    and return a retriever + llm instance.
    """
    embeddings = OpenAIEmbeddings(model=EMBEDDING_MODEL)

    if not os.path.isdir(faiss_dir):
        raise RuntimeError(f"FAISS directory not found: {faiss_dir}")

    vectorstore = LC_FAISS.load_local(faiss_dir, embeddings, allow_dangerous_deserialization=True)
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})

    llm = ChatOpenAI(model=LLM_MODEL_NAME, temperature=0)

    return retriever, llm

# ---------- Rule-based quick checks ----------
def rule_based_checks_doc(full_text: str) -> List[Dict]:
    issues = []
    t = full_text.lower()
    if "jurisdiction" not in t and "adgm" not in t:
        issues.append({"section": "Document-level", "issue": "Missing jurisdiction clause specifying ADGM", "severity": "High", "suggestion": "Add an explicit jurisdiction clause referencing ADGM Courts."})
    if not re.search(r"(signat|signed by|signature|for and on behalf)", t):
        issues.append({"section": "Document-level", "issue": "No signature block detected", "severity": "High", "suggestion": "Ensure a signature block (name, role, date) is present for all signatories."})
    if not re.search(r"\b(date[:\s]|[0-3]?\d\W+(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b|\b\d{4}\b)", t, re.IGNORECASE):
        issues.append({"section": "Document-level", "issue": "No clear date found", "severity": "Medium", "suggestion": "Add an execution date for the document."})
    return issues

def rule_based_checks_para(para_text: str) -> List[Dict]:
    issues = []
    t = para_text.lower()
    if re.search(r"(uae federal|federal courts|federal court|u.a.e\.)", t):
        issues.append({"section": "Paragraph", "issue": "References UAE Federal Courts instead of ADGM", "severity": "High", "suggestion": "Replace jurisdiction reference with ADGM Courts."})
    if re.search(r"\b(may|might|could|subject to|endeavour|best endeavours|reasonable endeavours)\b", t):
        issues.append({"section": "Paragraph", "issue": "Ambiguous or non-binding language (modal verbs)", "severity": "Medium", "suggestion": "Use clear, binding obligations where appropriate (e.g., 'shall' instead of 'may')."})
    if re.search(r"(executed by|in witness|for and on behalf|signed)\b", t) and not re.search(r"(signature|signed|name:)", t):
        issues.append({"section": "Paragraph", "issue": "Signatory/execution clause appears incomplete", "severity": "High", "suggestion": "Include signature, printed name, and capacity (e.g., director)."})
    if "indemn" in t and "liabilit" not in t:
        issues.append({"section": "Paragraph", "issue": "Indemnity clause present but lacks clarity on scope/limits", "severity": "Medium", "suggestion": "State scope, exceptions, and limits clearly."})
    return issues

# ---------- RAG-backed paragraph review ----------

def rag_check_paragraph(para_text: str, retriever, llm, k: int = 3) -> List[Dict]:
    try:
        retrieved = retriever.invoke(para_text)
        retrieved_docs = retrieved if isinstance(retrieved, list) else list(retrieved)
    except Exception:
        try:
            retrieved_docs = retriever.get_relevant_documents(para_text)
        except Exception:
            retrieved_docs = []

    evidence = []
    for d in retrieved_docs[:k]:
        m = getattr(d, "metadata", {}) or {}
        title = m.get("source_name") or m.get("source_url") or "ADGM reference"
        snippet = (d.page_content[:600] + "...") if len(d.page_content) > 600 else d.page_content
        evidence.append(f"Source: {title}\nExcerpt: {snippet}")

    prompt = f"""
You are an ADGM compliance reviewer. Analyze the paragraph below for ADGM-specific compliance issues
(invalid/missing clauses, wrong jurisdiction, ambiguous language, missing signatures/dates, or template mismatch).
Return ONLY valid JSON (an array). Each element must be an object with keys:
 - "section": "<clause or 'N/A'>",
 - "issue": "<brief issue>",
 - "severity": "High|Medium|Low",
 - "suggestion": "<fix>",
 - "citation": "<reference title or excerpt if applicable>"

Paragraph:
\"\"\"{para_text}\"\"\"

Evidence:
\"\"\"{chr(10).join(evidence)}\"\"\"
If there are no issues, return: []
"""
    resp = llm.invoke([HumanMessage(content=prompt)])
    raw = None
    if isinstance(resp, dict):
        raw = resp.get("content") or str(resp)
    else:
        raw = getattr(resp, "content", None) or str(resp)
    raw = raw.strip()

    try:
        parsed = json.loads(raw)
        if isinstance(parsed, list):
            return parsed
    except Exception:
        m = re.search(r"(\[.*\])", raw, re.S)
        if m:
            try:
                return json.loads(m.group(1))
            except:
                pass
    return []

# ---------- New annotation function ----------

def insert_comments_into_docx(
    input_docx_path: str,
    issues: List[Dict],
    output_docx_path: str
) -> str:
    """
    Inserts comments (inline brackets) into the docx paragraphs based on issues
    and appends a summary comments page at the end.
    
    Args:
        input_docx_path: Path to the input .docx file
        issues: List of issue dicts with keys like 'para_idx', 'issue', 'severity', 'suggestion', 'citation'
        output_docx_path: Path to save the annotated .docx file
    
    Returns:
        output_docx_path
    """
    doc = Document(input_docx_path)

    # group issues by paragraph index
    issues_by_para = {}
    for idx, issue in enumerate(issues, start=1):
        para_idx = issue.get("para_idx")
        if para_idx is not None:
            issues_by_para.setdefault(para_idx, []).append((idx, issue))

    # insert inline comments after the paragraph text
    for para_idx, comments in issues_by_para.items():
        if para_idx < len(doc.paragraphs):
            para = doc.paragraphs[para_idx]
            for comment_id, issue in comments:
                comment_text = f"  [COMMENT #{comment_id}: {issue.get('issue')}]"
            
                para.add_run("\n" + comment_text)

    # append a page break and add a summary page with all comments
    doc.add_page_break()
    doc.add_heading("Auto-generated Comments (Corporate Agent)", level=2)
    
    for idx, issue in enumerate(issues, start=1):
        doc.add_paragraph(f"Comment #{idx}")
        doc.add_paragraph(f"Document: {os.path.basename(input_docx_path)}")
        doc.add_paragraph(f"Paragraph idx: {issue.get('para_idx')}")
        doc.add_paragraph(f"Issue: {issue.get('issue')}")
        doc.add_paragraph(f"Severity: {issue.get('severity', 'Medium')}")
        if issue.get('suggestion'):
            doc.add_paragraph(f"Suggestion: {issue.get('suggestion')}")
        if issue.get('citation'):
            doc.add_paragraph(f"Citation: {issue.get('citation')}")
        doc.add_paragraph("")  # blank line for spacing

    doc.save(output_docx_path)
    return output_docx_path

# ---------- Main document review function ----------

def review_single_doc_and_save(path: str, doc_type: str, retriever, llm, output_dir: str = REPORTS_DIR,
                               para_rag_k: int = 2, max_paras: int = 50) -> Tuple[str, str, Dict]:
    paras = []
    doc_obj = Document(path)
    for idx, p in enumerate(doc_obj.paragraphs):
        txt = p.text.strip()
        if txt:
            paras.append({"para_idx": idx, "text": txt})
    full_text = "\n\n".join([p["text"] for p in paras])
    issues = []

    # document-level rule checks
    for it in rule_based_checks_doc(full_text):
        it.update({"document": os.path.basename(path), "para_idx": None})
        issues.append(it)

    # paragraph-level rules and RAG selection
    flagged = []
    kws = ["jurisdiction", "signat", "ubo", "ultimate beneficial", "adgm", "arbitrat", "indemn", "liabilit", "govern", "director", "member", "share"]
    for p in paras:
        p_issues = rule_based_checks_para(p["text"])
        if p_issues:
            for it in p_issues:
                it.update({"document": os.path.basename(path), "para_idx": p["para_idx"]})
                issues.append(it)
        if any(k in p["text"].lower() for k in kws) or len(p["text"]) > 300:
            flagged.append(p)
    flagged = flagged[:max_paras]

    # RAG-check flagged paragraphs
    for p in flagged:
        rag_issues = rag_check_paragraph(p["text"], retriever, llm, k=para_rag_k)
        for it in rag_issues:
            it.update({"document": os.path.basename(path), "para_idx": p["para_idx"]})
            issues.append(it)
    print("Unique paragraph indexes in issues:", set(it.get("para_idx") for it in issues))
    # annotate docx and save
    out_name = os.path.splitext(os.path.basename(path))[0] + "_reviewed.docx"
    out_path = os.path.join(output_dir, out_name)
    insert_comments_into_docx(path, issues, out_path)

    # save JSON report
    report = {"document": os.path.basename(path), "doc_type": doc_type, "issues": issues}
    json_path = os.path.join(output_dir, os.path.splitext(os.path.basename(path))[0] + "_report.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2, ensure_ascii=False)

    return out_path, json_path, report

# ----------wrapper used by Streamlit ----------

def process_documents(file_paths: List[str], faiss_dir: str = FAISS_DIR) -> Dict:
    """
    file_paths: list of local .docx file paths (already saved on disk)
    returns: structured dict with process, checklist, reports, reviewed file paths
    """
    retriever, llm = load_retriever_and_llm(faiss_dir)

    # classify uploaded
    classified = []
    for p in file_paths:
        txt = extract_docx_text(p)
        dtype = classify_doc_by_keywords(p.split(os.sep)[-1], txt)
        classified.append({"path": p, "doc_type": dtype, "text": txt})

    uploaded_types = [c["doc_type"] for c in classified]
    proc_name, missing = detect_process(uploaded_types)
    checklist = {
        "process": proc_name,
        "documents_uploaded": len(uploaded_types),
        "required_documents": len(ADGM_CHECKLISTS.get(proc_name, [])),
        "missing_documents": missing if missing else None
    }

    # run review for each document
    reports = []
    reviewed_files = []
    for c in classified:
        out_doc, out_json, rep = review_single_doc_and_save(c["path"], c["doc_type"], retriever, llm, output_dir=REPORTS_DIR)
        reports.append(rep)
        reviewed_files.append(out_doc)

    aggregate = {
        "process": proc_name,
        "checklist": checklist,
        "reports": reports,
        "reviewed_files": reviewed_files
    }
    # write aggregate
    agg_path = os.path.join(REPORTS_DIR, "aggregate_report.json")
    with open(agg_path, "w", encoding="utf-8") as f:
        json.dump(aggregate, f, indent=2, ensure_ascii=False)
    return aggregate